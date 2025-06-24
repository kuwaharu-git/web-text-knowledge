"""
Difyファイル生成モジュール
"""

import logging
import os
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime

from ..parser.text_parser import ParsedPage
from ..config.settings import Settings
from ..utils.logger import sanitize_filename, format_bytes


class DifyFileGenerator:
    """Difyファイル生成クラス"""
    
    def __init__(self, settings: Settings):
        self.settings = settings
        self.logger = logging.getLogger(__name__)
        
        # 統計情報
        self.stats = {
            "generated_files": [],
            "total_files": 0,
            "total_size": 0,
            "format_counts": {"txt": 0, "md": 0, "docx": 0}
        }
    
    def generate(self, pages: List[ParsedPage]) -> List[Path]:
        """Difyファイルを生成"""
        if not pages:
            self.logger.warning("生成するページがありません")
            return []
        
        self.logger.info(f"Difyファイル生成開始: {len(pages)} ページ")
        
        # ファイル分割が必要かチェック
        chunks = self._split_pages_into_chunks(pages)
        
        generated_files = []
        
        for i, chunk in enumerate(chunks):
            chunk_suffix = f"_part{i+1}" if len(chunks) > 1 else ""
            
            if self.settings.output_format in ["txt", "all"]:
                txt_file = self._generate_txt_file(chunk, chunk_suffix)
                if txt_file:
                    generated_files.append(txt_file)
            
            if self.settings.output_format in ["md", "all"]:
                md_file = self._generate_md_file(chunk, chunk_suffix)
                if md_file:
                    generated_files.append(md_file)
            
            if self.settings.output_format in ["docx", "all"]:
                docx_file = self._generate_docx_file(chunk, chunk_suffix)
                if docx_file:
                    generated_files.append(docx_file)
        
        self._log_stats()
        return generated_files
    
    def _split_pages_into_chunks(self, pages: List[ParsedPage]) -> List[List[ParsedPage]]:
        """ページをファイルサイズ制限に基づいてチャンクに分割"""
        chunks = []
        current_chunk = []
        current_size = 0
        max_size = self.settings.get_max_file_size_bytes()
        
        # ヘッダー情報のサイズを推定
        header_size = self._estimate_header_size()
        
        for page in pages:
            # ページのサイズを推定（UTF-8での文字数×3バイト + フォーマット用余白）
            page_size = len(page.content.encode('utf-8')) + len(page.title.encode('utf-8')) + 200
            
            # チャンクサイズが制限を超える場合
            if current_size + page_size + header_size > max_size and current_chunk:
                chunks.append(current_chunk)
                current_chunk = [page]
                current_size = page_size
            else:
                current_chunk.append(page)
                current_size += page_size
        
        if current_chunk:
            chunks.append(current_chunk)
        
        if len(chunks) > 1:
            self.logger.info(f"ファイルサイズ制限により {len(chunks)} 個のファイルに分割します")
        
        return chunks
    
    def _estimate_header_size(self) -> int:
        """ヘッダー情報のサイズを推定"""
        # メタデータ、区切り文字などのサイズを推定
        return 1000  # 1KB程度
    
    def _generate_txt_file(self, pages: List[ParsedPage], suffix: str = "") -> Optional[Path]:
        """テキストファイルを生成"""
        try:
            filename = self.settings.get_output_filename("txt")
            if suffix:
                filename = filename.replace(".txt", f"{suffix}.txt")
            
            file_path = self.settings.output_dir / sanitize_filename(filename)
            
            with open(file_path, 'w', encoding='utf-8') as f:
                # ヘッダー情報
                f.write(self._create_txt_header(pages))
                f.write("\n")
                
                # ページ内容
                for i, page in enumerate(pages):
                    if i > 0:
                        f.write("\n" + "="*50 + "\n\n")
                    
                    f.write(self._format_page_as_txt(page))
                
                # フッター情報
                f.write("\n\n" + "="*50 + "\n")
                f.write(self._create_txt_footer(pages))
            
            file_size = file_path.stat().st_size
            self.stats["generated_files"].append(str(file_path))
            self.stats["total_size"] += file_size
            self.stats["format_counts"]["txt"] += 1
            
            self.logger.info(f"TXTファイル生成完了: {file_path} ({format_bytes(file_size)})")
            return file_path
        
        except Exception as e:
            self.logger.error(f"TXTファイル生成エラー: {e}")
            return None
    
    def _generate_md_file(self, pages: List[ParsedPage], suffix: str = "") -> Optional[Path]:
        """Markdownファイルを生成"""
        try:
            filename = self.settings.get_output_filename("md")
            if suffix:
                filename = filename.replace(".md", f"{suffix}.md")
            
            file_path = self.settings.output_dir / sanitize_filename(filename)
            
            with open(file_path, 'w', encoding='utf-8') as f:
                # ヘッダー情報
                f.write(self._create_md_header(pages))
                f.write("\n\n")
                
                # 目次（オプション）
                if self.settings.output.add_table_of_contents and len(pages) > 5:
                    f.write(self._create_md_toc(pages))
                    f.write("\n\n")
                
                # ページ内容
                for i, page in enumerate(pages):
                    if i > 0:
                        f.write("\n---\n\n")
                    
                    f.write(self._format_page_as_md(page))
                
                # フッター情報
                f.write("\n\n---\n\n")
                f.write(self._create_md_footer(pages))
            
            file_size = file_path.stat().st_size
            self.stats["generated_files"].append(str(file_path))
            self.stats["total_size"] += file_size
            self.stats["format_counts"]["md"] += 1
            
            self.logger.info(f"MDファイル生成完了: {file_path} ({format_bytes(file_size)})")
            return file_path
        
        except Exception as e:
            self.logger.error(f"MDファイル生成エラー: {e}")
            return None
    
    def _generate_docx_file(self, pages: List[ParsedPage], suffix: str = "") -> Optional[Path]:
        """Word文書ファイルを生成"""
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.style import WD_STYLE_TYPE
            
            filename = self.settings.get_output_filename("docx")
            if suffix:
                filename = filename.replace(".docx", f"{suffix}.docx")
            
            file_path = self.settings.output_dir / sanitize_filename(filename)
            
            doc = Document()
            
            # ドキュメントのスタイル設定
            self._setup_docx_styles(doc)
            
            # ヘッダー情報
            header_info = self._create_docx_header(pages)
            title_paragraph = doc.add_paragraph()
            title_run = title_paragraph.add_run(header_info["title"])
            title_run.bold = True
            title_paragraph.style = "Title"
            
            # メタデータ
            for key, value in header_info["metadata"].items():
                para = doc.add_paragraph()
                para.add_run(f"{key}: ").bold = True
                para.add_run(str(value))
            
            doc.add_page_break()
            
            # 目次（オプション）
            if self.settings.output.add_table_of_contents and len(pages) > 5:
                toc_heading = doc.add_paragraph()
                toc_heading.add_run("目次").bold = True
                toc_heading.style = "Heading 1"
                
                for i, page in enumerate(pages, 1):
                    toc_item = doc.add_paragraph()
                    toc_item.add_run(f"{i}. {page.title or 'Untitled'}")
                    toc_item.style = "List Number"
                
                doc.add_page_break()
            
            # ページ内容
            for i, page in enumerate(pages):
                if i > 0:
                    doc.add_page_break()
                
                self._add_page_to_docx(doc, page)
            
            # フッター情報
            doc.add_page_break()
            footer_paragraph = doc.add_paragraph()
            footer_run = footer_paragraph.add_run("生成情報")
            footer_run.bold = True
            footer_paragraph.style = "Heading 1"
            
            footer_info = self._create_docx_footer(pages)
            for key, value in footer_info.items():
                para = doc.add_paragraph()
                para.add_run(f"{key}: ").bold = True
                para.add_run(str(value))
            
            doc.save(file_path)
            
            file_size = file_path.stat().st_size
            self.stats["generated_files"].append(str(file_path))
            self.stats["total_size"] += file_size
            self.stats["format_counts"]["docx"] += 1
            
            self.logger.info(f"DOCXファイル生成完了: {file_path} ({format_bytes(file_size)})")
            return file_path
        
        except ImportError:
            self.logger.error("python-docxライブラリがインストールされていません")
            return None
        except Exception as e:
            self.logger.error(f"DOCXファイル生成エラー: {e}")
            return None
    
    def _create_txt_header(self, pages: List[ParsedPage]) -> str:
        """TXTファイルのヘッダーを作成"""
        total_chars = sum(len(page.content) for page in pages)
        total_tokens = sum(page.tokens for page in pages)
        
        header = f"""サイト名: {self.settings.url}
取得日時: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
総ページ数: {len(pages)}
総文字数: {total_chars:,}
総トークン数: {total_tokens:,}
{"="*50}"""
        return header
    
    def _create_txt_footer(self, pages: List[ParsedPage]) -> str:
        """TXTファイルのフッターを作成"""
        footer = f"""
生成完了日時: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
生成システム: Difyナレッジファイル生成システム v2.0
注意: このファイルはWebスクレイピングにより自動生成されました。
元サイトの利用規約を遵守してご利用ください。"""
        return footer
    
    def _format_page_as_txt(self, page: ParsedPage) -> str:
        """ページをTXT形式でフォーマット"""
        fetch_time = datetime.fromisoformat(page.metadata["fetch_time"]).strftime('%Y-%m-%d %H:%M:%S')
        
        formatted = f"""【{page.title}】
URL: {page.url}
取得日時: {fetch_time}
文字数: {page.metadata.get('character_count', 0):,}
トークン数: {page.tokens:,}

{page.content}"""
        
        return formatted
    
    def _create_md_header(self, pages: List[ParsedPage]) -> str:
        """Markdownファイルのヘッダーを作成"""
        total_chars = sum(len(page.content) for page in pages)
        total_tokens = sum(page.tokens for page in pages)
        
        header = f"""# {self.settings.site_name.title()} - ナレッジベース

**取得元URL**: {self.settings.url}  
**取得日時**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  
**総ページ数**: {len(pages)}  
**総文字数**: {total_chars:,}  
**総トークン数**: {total_tokens:,}"""
        
        return header
    
    def _create_md_toc(self, pages: List[ParsedPage]) -> str:
        """Markdownの目次を作成"""
        toc = "## 目次\n\n"
        for i, page in enumerate(pages, 1):
            title = page.title or "Untitled"
            # Markdownのリンク形式に変換
            anchor = title.lower().replace(" ", "-").replace("　", "-")
            toc += f"{i}. [{title}](#{anchor})\n"
        
        return toc
    
    def _create_md_footer(self, pages: List[ParsedPage]) -> str:
        """Markdownファイルのフッターを作成"""
        footer = f"""## 生成情報

**生成完了日時**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  
**生成システム**: Difyナレッジファイル生成システム v2.0  
**注意**: このファイルはWebスクレイピングにより自動生成されました。元サイトの利用規約を遵守してご利用ください。"""
        
        return footer
    
    def _format_page_as_md(self, page: ParsedPage) -> str:
        """ページをMarkdown形式でフォーマット"""
        fetch_time = datetime.fromisoformat(page.metadata["fetch_time"]).strftime('%Y-%m-%d %H:%M:%S')
        
        # タイトルをMarkdownのヘッダーに
        title = page.title or "Untitled"
        
        formatted = f"""## {title}

**URL**: {page.url}  
**取得日時**: {fetch_time}  
**文字数**: {page.metadata.get('character_count', 0):,}  
**トークン数**: {page.tokens:,}

{page.content}"""
        
        return formatted
    
    def _setup_docx_styles(self, doc):
        """Word文書のスタイルを設定"""
        try:
            from docx.shared import Pt
            from docx.enum.style import WD_STYLE_TYPE
            
            # 基本スタイルの設定（エラーが発生しても続行）
            styles = doc.styles
            
            # タイトルスタイル
            if 'Title' not in [style.name for style in styles]:
                title_style = styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
                title_style.font.size = Pt(18)
                title_style.font.bold = True
        except:
            # スタイル設定に失敗しても続行
            pass
    
    def _create_docx_header(self, pages: List[ParsedPage]) -> Dict[str, Any]:
        """Word文書のヘッダー情報を作成"""
        total_chars = sum(len(page.content) for page in pages)
        total_tokens = sum(page.tokens for page in pages)
        
        return {
            "title": f"{self.settings.site_name.title()} - ナレッジベース",
            "metadata": {
                "取得元URL": self.settings.url,
                "取得日時": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                "総ページ数": len(pages),
                "総文字数": f"{total_chars:,}",
                "総トークン数": f"{total_tokens:,}"
            }
        }
    
    def _create_docx_footer(self, pages: List[ParsedPage]) -> Dict[str, Any]:
        """Word文書のフッター情報を作成"""
        return {
            "生成完了日時": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            "生成システム": "Difyナレッジファイル生成システム v2.0",
            "注意": "このファイルはWebスクレイピングにより自動生成されました。元サイトの利用規約を遵守してご利用ください。"
        }
    
    def _add_page_to_docx(self, doc, page: ParsedPage):
        """Word文書にページを追加"""
        try:
            from docx.shared import Pt
            
            # ページタイトル
            title_paragraph = doc.add_paragraph()
            title_run = title_paragraph.add_run(page.title or "Untitled")
            title_run.bold = True
            title_run.font.size = Pt(14)
            
            # メタデータ
            fetch_time = datetime.fromisoformat(page.metadata["fetch_time"]).strftime('%Y-%m-%d %H:%M:%S')
            
            metadata_paragraph = doc.add_paragraph()
            metadata_paragraph.add_run("URL: ").bold = True
            metadata_paragraph.add_run(page.url)
            
            metadata_paragraph = doc.add_paragraph()
            metadata_paragraph.add_run("取得日時: ").bold = True
            metadata_paragraph.add_run(fetch_time)
            
            metadata_paragraph = doc.add_paragraph()
            metadata_paragraph.add_run("文字数: ").bold = True
            metadata_paragraph.add_run(f"{page.metadata.get('character_count', 0):,}")
            
            metadata_paragraph = doc.add_paragraph()
            metadata_paragraph.add_run("トークン数: ").bold = True
            metadata_paragraph.add_run(f"{page.tokens:,}")
            
            # 空行
            doc.add_paragraph()
            
            # コンテンツ
            content_paragraphs = page.content.split('\n\n')
            for para_text in content_paragraphs:
                if para_text.strip():
                    doc.add_paragraph(para_text.strip())
        
        except Exception as e:
            self.logger.warning(f"Word文書へのページ追加でエラー: {e}")
            # エラーが発生してもテキストのみ追加
            doc.add_paragraph(page.title or "Untitled")
            doc.add_paragraph(page.content)
    
    def _log_stats(self):
        """統計情報をログ出力"""
        self.logger.info("=== ファイル生成統計 ===")
        self.logger.info(f"生成ファイル数: {len(self.stats['generated_files'])}")
        self.logger.info(f"総ファイルサイズ: {format_bytes(self.stats['total_size'])}")
        
        for format_type, count in self.stats["format_counts"].items():
            if count > 0:
                self.logger.info(f"{format_type.upper()}ファイル: {count}個")
        
        for file_path in self.stats["generated_files"]:
            self.logger.info(f"  - {file_path}")
    
    def get_stats(self) -> Dict[str, Any]:
        """統計情報を取得"""
        return self.stats.copy()
